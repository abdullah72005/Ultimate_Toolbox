import os
import pandas as pd
from flask import send_file
from PIL import Image
from pydub import AudioSegment
from werkzeug.utils import secure_filename
import fitz
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

from helpers.functions import deleteFiles

uploadFolder = 'static/uploads/'

# link magic file type input with actual file type input via dictionary
imagetypes = {
    'image/jpeg': 'jpeg',
    'image/png': 'png',
    'image/tiff': 'tiff',
    'image/gif': 'gif',
    'image/bmp': 'bmp',
    'image/vnd.microsoft.icon': 'ico',
    'image/x-pcx': 'pcx',
    'image/x-portable-pixmap': 'ppm',
    'image/vnd.adobe.photoshop': 'psd',
    'image/webp': 'webp'
}

audioTypes = {
    'audio/x-wav': 'wav',
    'audio/mpeg': 'mp3',
    'audio/x-hx-aac-adts': 'aac',
    'audio/x-m4a': 'm4a',
    'audio/vnd.dolby.dd-raw': 'ac3',
    'audio/amr': 'amr'

}
  
txttypes = {
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
    'application/pdf': 'pdf',
    'text/plain': 'txt'
}

csvtypes = {
    'text/csv': 'csv',
    'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet':'xlsx',
    'application/json': 'json'
}

def conImgtoImg(fileName, file, app, choice):
    # Construct the path to the input image file
    img_path = os.path.join(uploadFolder, file)

    # Open the image file using Pillow (PIL)
    img = Image.open(img_path)

    # Construct the path to the output PNG file
    png_path = os.path.join(uploadFolder, f"{fileName}.{choice}")
    choice2 = choice.upper()

    # Save the image as a PNG file
    img.save(png_path, choice2)

def getOutputChoices(extension, ime, txte, aue, csve, im, txt, au, csv):
    # if file is image let output choices be image options and remove the file extention
    if extension in ime:
        x = [choice for choice in im if choice != imagetypes[extension] and choice not in ['pcx', 'psd']] + ['pdf']
        return x


    # if file is txt let output choices be txt options and remove the file extention
    if extension in txte:
        x = [choice for choice in txt if choice != txttypes[extension]]
        return x
    
    # if file is audio let output choices be audio options and remove the file extention
    if extension in aue:
        notto = ["aac", 'm4a', "amr"]
        x = [choice for choice in au if choice != audioTypes[extension] and choice not in notto]
        return x
    
    # if file is csv let output choices be csv options and remove the file extention
    if extension in csve:
        x = [choice for choice in csv if choice != csvtypes[extension]]
        return x 


def convIMAGE(fileName, file, app, choice):
    # Check the user's choice and perform the corresponding conversion
    
    # Convert image to PDF
    conImgtoImg(fileName, file, app, choice)
    pdf_path = os.path.join(uploadFolder, f"{fileName}.{choice}"
    )  

    # make variable with output file, delete files from local directory, return output file
    outputFile = send_file(pdf_path, as_attachment=True)
    deleteFiles(uploadFolder)
    return outputFile

def word2pdf(file, fileName, choice):
    # Combine folder path with the uploaded file name
    print("It is in the function")
    word_path = os.path.join(uploadFolder, file)
    pdf_path = os.path.join(uploadFolder, f"{fileName}.{choice}") 

    # Read content from the DOCX file
    doc = Document(word_path)

    # Extract text from paragraphs and add newline between them
    print("I am here")
    text_content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])

    # Create a PDF using reportlab
    pdf_canvas = canvas.Canvas(pdf_path, pagesize=letter)
    width, height = letter

    # Set font and size
    pdf_canvas.setFont("Helvetica", 12)
    print("I am there")
    # Split content by lines to handle newline characters
    lines = text_content.split('\n')
    print(lines)

    # Write content to PDF
    for i, line in enumerate(lines):
        y_position = height - 50 - (i * 14)  # Adjust spacing

        # Strip leading spaces and then draw the line
        pdf_canvas.drawString(50, y_position, line.lstrip())

    # Save the PDF
    print("I am everywhere")
    pdf_canvas.save()
    # Send the Word file as an attachment and delete temporary files
    outputFile = send_file(pdf_path, as_attachment=True)
    deleteFiles(uploadFolder)

    return outputFile

def pdf2word(file, fileName, choice):
    # Combine folder path with the uploaded file name
    pdf_path = os.path.join(uploadFolder, file)

    # Open the PDF file using PyMuPDF
    doc = fitz.open(pdf_path)

    # Create a new Word document using python-docx
    document = Document()

    # Loop through each page in the PDF
    for page in doc:

        # Extract text from the page
        text = page.get_text("text")

        # Add each paragraph as a new paragraph in the Word document
        for paragraph in text.split('\n'):
            document.add_paragraph(paragraph)

    # Combine folder path with the output Word file name
    word_path = os.path.join(uploadFolder, f"{fileName}.{choice}")  
    document.save(word_path)

    # Send the Word file as an attachment and delete temporary files
    outputFile = send_file(word_path, as_attachment=True)
    deleteFiles(uploadFolder)

    # Return the Word file   
    return outputFile


def txt2word(file, fileName, choice):
    # Combine folder path with the uploaded file name
    txt_path = os.path.join(uploadFolder, file)

    # Read the content from the text file
    with open(txt_path, 'r', encoding='utf-8') as file:
        text_content = file.read()

    # Create a new Word document using python-docx
    doc = Document()

    # Add the text content to the Word document
    doc.add_paragraph(text_content)

    # Combine folder path with the output Word file name
    docx_path = os.path.join(uploadFolder, f"{fileName}.{choice}")  

    # Save the Word document
    doc.save(docx_path)

    # Send the Word file as an attachment and delete temporary files
    outputFile = send_file(docx_path, as_attachment=True)
    deleteFiles(uploadFolder)
    return outputFile


def txt2pdf(file, fileName, choice):
    # Combine folder path with the uploaded file name
    txt_path = os.path.join(uploadFolder, file)
    pdf_path = os.path.join(uploadFolder, f"{fileName}.{choice}")  

    # Read text from the input file
    with open(txt_path, 'r') as file:
        content = file.read()

    # Replace newline characters with proper formatting
    content = content.replace('\n', '<br/>')

    # Replace special characters with spaces using Unicode replacement
    content = content.encode('utf-8', 'replace').decode('utf-8')

    # Create a PDF
    pdf_canvas = canvas.Canvas(pdf_path, pagesize=letter)
    width, height = letter

    # Set font and size
    pdf_canvas.setFont("Helvetica", 12)

    # Split content by <br/> to handle newline characters
    lines = content.split('<br/>')

    # Write content to PDF
    for i, line in enumerate(lines):
        y_position = height - 50 - (i * 14)  # Adjust spacing

        # Strip leading spaces and then draw the line
        pdf_canvas.drawString(50, y_position, line.lstrip())

    # Save the PDF
    pdf_canvas.save()

    # Send the PDF file as an attachment and delete temporary files
    outputFile = send_file(pdf_path, as_attachment=True)
    deleteFiles(uploadFolder)
    return outputFile


def word2txt(file, fileName, choice):

    # Combine folder path with the uploaded file name
    docx_path = os.path.join(uploadFolder, file)

    # Combine folder path with the output TXT file name
    txt_path = os.path.join(uploadFolder, f"{fileName}.{choice}")  

    # Open the Word document using python-docx
    doc = Document(docx_path)

    # Extract text from paragraphs and add newline between them
    text_content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])

    # Write the text content to a TXT file
    with open(txt_path, 'w') as txt_file:
        txt_file.write(text_content)
    
    # Send the TXT file as an attachment and delete temporary files
    outputFile = send_file(txt_path, as_attachment=True)
    deleteFiles(uploadFolder)
    return outputFile


def pdf2txt(file, fileName, choice):   
    # Combine folder path with the uploaded file name 
    pdf_path = os.path.join(uploadFolder, file)

    # Combine folder path with the output TXT file name
    txt_path = os.path.join(uploadFolder, f"{fileName}.{choice}")

    # Open the PDF file
    pdf_document = fitz.open(pdf_path)

    text_content = ""

    # Iterate through pages and extract text
    for page_number in range(pdf_document.page_count):
        page = pdf_document.load_page(page_number)
        text_content += page.get_text()

    # Write the text content to a TXT file
    with open(txt_path, 'w', encoding='utf-8') as txt_file:
        txt_file.write(text_content)
    
    # Send the TXT file as an attachment and delete temporary files
    outputFile = send_file(txt_path, as_attachment=True)
    deleteFiles(uploadFolder)
    return outputFile


def convert_audio(input_file, extension, filename, choice):
    # get the audio path and Load the input audio file
    audio_path = os.path.join(uploadFolder, input_file)
    sound = AudioSegment.from_file(audio_path, format=audioTypes[extension])
    
    # get output filename and get its path
    output_filename = secure_filename(filename + '.' + choice)
    output_filepath = os.path.join(uploadFolder, output_filename)

    # Export the sound to the output format, delete files, and return the output file
    sound.export(output_filepath, format=choice)
    x = send_file(output_filepath, as_attachment=True)
    deleteFiles(uploadFolder)
    return x


def convert_csv(input_file, extension, filename, choice):
    input_path = os.path.join(uploadFolder, input_file)
    filetype = csvtypes[extension]
    output_filename = filename + '.' + choice
    output_path = os.path.join(uploadFolder, output_filename)

    # if input file is a csv file
    if filetype == 'csv':
        print('its a csv')

        # if choice is excel 
        if choice == 'xlsx':
            
            csv = pd.read_csv(input_path, encoding = "ISO-8859-1")
            ew = pd.ExcelWriter(output_path)
            csv.to_excel(ew)
            ew.close()
            x = send_file(output_path, as_attachment=True)
        
        # if he chooses json
        elif choice == 'json':

            csv = pd.read_csv(input_path, encoding = "ISO-8859-1")
            csv.to_json(output_path, indent = 1, orient = 'records')
            x = send_file(output_path, as_attachment=True)

    # if input file is a excel file
    elif filetype == 'xlsx':
        
        # if user chooses csv
        if choice == 'csv':
            
            xl = pd.read_excel(input_path)
            xl.to_csv(output_path, index=None, header=True)
            x = send_file(output_path, as_attachment=True)

        # if user chooses json
        elif choice == 'json':

            xl = pd.read_excel(input_path)
            xl.to_json(output_path, indent = 1, orient = 'records')
            x = send_file(output_path, as_attachment=True)

    # if input file type is json
    elif filetype == 'json':

        # if user chooses csv
        if choice == 'csv':

            js = pd.read_json(input_path)
            js.to_csv(output_path, index=None, header=True)
            x = send_file(output_path, as_attachment=True)

        # if choice is excel 
        elif choice == 'xlsx':
            js = pd.read_json(input_path)
            ew = pd.ExcelWriter(output_path)
            js.to_excel(ew)
            ew.close()
            x = send_file(output_path, as_attachment=True)


    deleteFiles(uploadFolder)
    return x



