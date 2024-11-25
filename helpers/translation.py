from googletrans import Translator
import os
from docx import Document
from flask import send_file
from helpers.functions import deleteFiles

uploadFolder = 'static/uploads/'

def translatetxt(input_txt, input_lang, output_lang, langs):
    translator = Translator()

    # if no input lang selected, detect input lang then translate and return translated txt
    if input_lang == 'detect':
        g = translator.detect(input_txt)
        input_lang = g.lang
        output = translator.translate(input_txt, src=input_lang, dest=langs[output_lang])
    else:
        output = translator.translate(input_txt, src=langs[input_lang], dest=langs[output_lang])
    
    return output.text

def trans_doc(input_file, extension, fileName, input_lang, output_lang, langs):

    if extension == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' or extension == 'docx':
        # Load the input Word document
        doc = Document(input_file)

        # Initialize the translator
        translator = Translator()

        # Create a new Word document to store the translated content
        new_document = Document()

        docText = ""

        for paragraph in doc.paragraphs:
            docText += str(paragraph.text) + "\n"

        # Separate string into segmates to be able to translate it
        chunks = [docText[i:i + 15000] for i in range(0, len(docText), 15000)]
        if len(chunks) > 10:
            return 1

        
        # Determine the source language if it is set to 'detect'
        srcLang = input_lang
        
        # make sure the paragraph includes text to prevent errors
        if docText:
            
            if srcLang == 'detect':
                g = translator.detect(docText)
                srcLang = g.lang

                for chunk in chunks:
                    # Translate the paragraph from detect language to the output language
                    output = translator.translate(chunk, src=srcLang, dest=langs[output_lang]).text

                    # Add the translated paragraph to the new document
                    new_document.add_paragraph(output)

            else:
                for chunk in chunks:
                    # Translate the paragraph from the specified input language to the output language
                    output = translator.translate(chunk, src=langs[input_lang], dest=langs[output_lang]).text

                    # Add the translated paragraph to the new document
                    new_document.add_paragraph(output)

        # Save the translated document with a new filename
        new_filename = f"{fileName}-{output_lang}.docx"
        new_document.save(os.path.join(uploadFolder, new_filename))

        # Create a Flask send_file response for downloading the translated document
        word_path = os.path.join(uploadFolder, new_filename)
        outputFile = send_file(word_path, as_attachment=True)

    
    elif extension == 'text/plain' or extension == 'txt':

        # declare input and output paths
        new_filename = f"{fileName}-{output_lang}.txt"
        output_filepath = os.path.join(uploadFolder, new_filename)
        input_filepath = os.path.join(uploadFolder, f"{fileName}.txt")

        # open the file and read its content
        with open(input_filepath, 'r') as file:
            text = file.read()

        # Initialize the translator
        translator = Translator()
        
        # Determine the source language if it is set to 'detect'
        srcLang = input_lang
        
        # make sure that there is text to translate
        if text:
            if srcLang == 'detect':
                g = translator.detect(text)
                srcLang = g.lang

                # Translate the paragraph from detect language to the output language
                transtxt = str(translator.translate(text, src=srcLang, dest=langs[output_lang]).text)
            else:
                # Translate the paragraph from the specified input language to the output language
                transtxt = str(translator.translate(text, src=langs[input_lang], dest=langs[output_lang]).text)
        else:
            transtxt = None

        
        # Write the text content to a TXT file
        with open(output_filepath, 'w') as txt_file:
            txt_file.write(transtxt)

        # declare outputFile
        outputFile = send_file(output_filepath, as_attachment=True)


    # Clean up the upload folder after processing
    deleteFiles(uploadFolder)

    # Return the Flask send_file response
    return outputFile
